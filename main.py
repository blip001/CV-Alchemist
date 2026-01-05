import os
import uuid
import shutil
import logging
import tempfile
import json
import re
from fastapi import FastAPI, UploadFile, File, Form
from fastapi import Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.responses import RedirectResponse
from pydantic import BaseModel

from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, EmailStr
from fastapi_mail import ConnectionConfig, FastMail, MessageSchema, MessageType
from dotenv import load_dotenv

from google.cloud import secretmanager
import settings
import stripe
import vertexai
from google.auth import default
from vertexai.generative_models import GenerativeModel
from pdfminer.high_level import extract_text as extract_pdf
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

logging.basicConfig(level=logging.INFO)
PROJECT_ID = "cv-alchemist-482203"
LOCATION = "us-central1" 


app = FastAPI()
analysis_cache = {}
mail_config = None

def access_secret_version(secret_id: str, project_id: str = PROJECT_ID, version_id: str = "latest") -> str | None:
    """
    Access the payload for the given secret version and return it.
    Returns None if the secret cannot be accessed.
    """
    try:
        client = secretmanager.SecretManagerServiceClient()
        name = f"projects/{project_id}/secrets/{secret_id}/versions/{version_id}"
        response = client.access_secret_version(name=name)
        return response.payload.data.decode("UTF-8")
    except Exception as e:
        logging.warning(f"Could not access secret {secret_id}. Error: {e}. Falling back to environment variables.")
        return None

@app.on_event("startup")
def startup_event():
    global mail_config
    load_dotenv()
    # Load Stripe API Key from Secret Manager or environment variable
    stripe.api_key = access_secret_version("STRIPE_API_KEY") or os.getenv("STRIPE_API_KEY")
    if not stripe.api_key:
        logging.warning("Stripe API key not found in Secret Manager or environment variables.")

    # Load Mail Configuration once on startup
    try:
        mail_server = access_secret_version("MAIL_SERVER") or os.getenv("MAIL_SERVER")
        if not mail_server:
            raise ValueError("MAIL_SERVER secret is not set or accessible.")

        mail_config = ConnectionConfig(
            MAIL_USERNAME=access_secret_version("MAIL_USERNAME") or os.getenv("MAIL_USERNAME"),
            MAIL_PASSWORD=access_secret_version("MAIL_PASSWORD") or os.getenv("MAIL_PASSWORD"),
            MAIL_FROM=access_secret_version("MAIL_FROM") or os.getenv("MAIL_FROM"),
            MAIL_PORT=int(access_secret_version("MAIL_PORT") or os.getenv("MAIL_PORT", 587)),
            MAIL_SERVER=mail_server,
            MAIL_FROM_NAME=access_secret_version("MAIL_FROM_NAME") or os.getenv("MAIL_FROM_NAME"),
            MAIL_STARTTLS=(access_secret_version("MAIL_STARTTLS") or os.getenv("MAIL_STARTTLS", "True")).lower() == "true",
            MAIL_SSL_TLS=(access_secret_version("MAIL_SSL_TLS") or os.getenv("MAIL_SSL_TLS", "False")).lower() == "true",
            USE_CREDENTIALS=True,
            VALIDATE_CERTS=True,
            TEMPLATE_FOLDER=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
        )
        logging.info("Mail configuration loaded successfully.")
    except Exception as e:
        logging.error(f"Failed to load mail configuration on startup: {e}")

    # Initialize Vertex AI
    try:
        credentials, discovered_project_id = default()
        # Use discovered project ID if available, otherwise fall back to the hardcoded one.
        project_to_use = discovered_project_id or PROJECT_ID
        if not project_to_use:
             raise Exception("Google Cloud project ID is not set or could not be discovered.")
        vertexai.init(project=project_to_use, location=LOCATION, credentials=credentials)
        logging.info("Vertex AI initialized successfully.")
    except Exception as e:
        logging.error(f"Failed to initialize Vertex AI: {e}")


app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

@app.get("/")
async def read_root():
    # Robustly find index.html
    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, "index.html")
    if os.path.exists(file_path):
        return FileResponse(file_path, headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    return HTMLResponse('<h1>Error: index.html not found</h1><p><a href="/">Go to Homepage</a></p>', status_code=500)

@app.get("/result/{result_id}")
async def get_result(result_id: str):
    """Retrieve a cached analysis result."""
    result = analysis_cache.get(result_id)
    if not result:
        return JSONResponse(status_code=404, content={"detail": "Result not found or expired."})
    return result

@app.post("/analyze")
async def analyze(file: UploadFile = File(...), job_title: str = Form("")):
    text = ""
    temp_file_path = ""
    file_extension = os.path.splitext(file.filename)[1].lower()

    try:
        # Save uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            shutil.copyfileobj(file.file, temp_file)
            temp_file_path = temp_file.name

        # Extract text based on file type
        if file_extension == ".pdf":
            text = extract_pdf(temp_file_path)
        elif file_extension == ".docx":
            doc = docx.Document(temp_file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            return JSONResponse(status_code=400, content={"detail": "Unsupported file type. Please upload a PDF or DOCX."})

        if not text.strip():
            return JSONResponse(status_code=400, content={"detail": "Could not extract text from the document. It might be an image-only file."})

        # Call Vertex AI for analysis
        model = GenerativeModel("gemini-2.0-flash-lite-001")
        if job_title and job_title.strip():
            prompt = f"""Analyze the following resume for the position of '{job_title}'.
    Provide a score and feedback based on its suitability for that specific role.
    Resume Text: 
    {text[:8000]}
    
    Return your response as a JSON object with two keys: "score" (an integer between 0 and 100), and "feedback" (a list of short, actionable feedback strings).
    Example: {{"score": 85, "feedback": ["Great use of action verbs.", "Consider adding a summary section."]}}
    """
        else:
            prompt = f"""Analyze the following resume and provide a general analysis.
    Provide a score and feedback on its overall quality, structure, and clarity.
    Resume Text: 
    {text[:8000]}
    
    Return your response as a JSON object with two keys: "score" (an integer between 0 and 100), and "feedback" (a list of short, actionable feedback strings for general improvement).
    Example: {{"score": 75, "feedback": ["The resume is well-structured.", "Consider quantifying achievements with numbers."]}}
    """
        response = model.generate_content(prompt)
        
        # Clean and parse the JSON response
        response_text = response.text
        # Use regex to find the JSON object within the response text
        match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if match:
            json_str = match.group(0)
            try:
                analysis_result = json.loads(json_str)
            except json.JSONDecodeError:
                raise ValueError("Failed to decode JSON from the AI response.")
        else:
            raise ValueError("Could not find a valid JSON object in the AI response.")

        analysis_result["raw_text_preview"] = text[:2000]

        # Cache the result
        result_id = str(uuid.uuid4())
        analysis_cache[result_id] = analysis_result
        analysis_result["result_id"] = result_id
        return analysis_result 

    except Exception as e:
        logging.error(f"Failed to analyze file: {e}")
        return JSONResponse(status_code=500, content={"detail": f"An error occurred during analysis: {e}"})
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

@app.post("/rewrite")
async def rewrite(req: dict):
    try:
        job_title = req.get('job_title', 'the user-specified position')
        model = GenerativeModel("gemini-2.0-flash-lite-001")
        prompt = f"Rewrite this resume to be highly optimized for the position of '{job_title}'. Return text only:\n{req.get('text','')[:8000]}"
        res = model.generate_content(prompt)
        return {"optimized_text": res.text.replace("```", "")}
    except Exception as e:
        logging.error(f"Error during rewrite: {e}")
        return {"optimized_text": f"Error: {str(e)}"}


# --- Contact Form ---

current_dir = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=os.path.join(current_dir, "templates"))

@app.get("/contact", response_class=HTMLResponse)
async def get_contact_form(request: Request):
    """Displays the contact form page."""
    return templates.TemplateResponse("contact.html", {"request": request})

@app.post("/contact")
async def post_contact_form(
    request: Request,
    name: str = Form(...),
    email: EmailStr = Form(...),
    message: str = Form(...)
):
    """Handles contact form submission."""
    html = f"""
    <p>You have a new contact form submission from:</p>
    <ul>
        <li><b>Name:</b> {name}</li>
        <li><b>Email:</b> {email}</li>
    </ul>
    <p><b>Message:</b></p>
    <p>{message}</p>
    """

    message_to_send = MessageSchema(
        subject="New Contact Form Submission",
        recipients=["support@cvalchemist.com"],
        body=html,
        subtype=MessageType.html
    )

    if not mail_config:
        logging.error("Mail configuration is not available.")
        return RedirectResponse("/?contact_status=error#contact-section", status_code=303)

    fm = FastMail(mail_config)
    try:
        await fm.send_message(message_to_send)
        return RedirectResponse("/?contact_status=success#contact-section", status_code=303)
    except Exception as e:
        logging.error(f"Failed to send email: {e}")
        return RedirectResponse("/?contact_status=error#contact-section", status_code=303)

# --- End Contact Form ---

@app.post("/create-checkout-session")
async def checkout(req: dict):
    if not stripe.api_key:
        logging.error("Stripe API key is not configured.")
        return JSONResponse(status_code=500, content={"detail": "Payment processor is not configured."})
    try:
        s = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{'price_data': {'currency': 'usd', 'product_data': {'name': 'Resume Optimization'}, 'unit_amount': 2999}, 'quantity': 1}],
            mode='payment',
            success_url=f"{req.get('origin_url')}?success=true",
            cancel_url=f"{req.get('origin_url')}?canceled=true",
            allow_promotion_codes=True,
        )
        return {"url": s.url}
    except Exception as e:
        return JSONResponse(status_code=500, content={"detail": str(e)})

class DL(BaseModel): text: str
@app.post("/download-pdf")
async def dl_pdf(r: DL): 
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        f = temp_file.name
        c = canvas.Canvas(f, pagesize=letter)
        text_object = c.beginText(40, 750)
        text_object.setFont("Helvetica", 10)
        for line in r.text.split('\n'):
            text_object.textLine(line[:100])
        c.drawText(text_object)
        c.save()
    
    return FileResponse(f, filename="resume.pdf", background=lambda: os.remove(f))

@app.post("/download-docx")
async def dl_docx(r: DL):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        f = temp_file.name
        d = docx.Document()
        d.add_paragraph(r.text)
        d.save(f)
    
    return FileResponse(f, filename="resume.docx", background=lambda: os.remove(f))

if __name__ == "__main__":
    import uvicorn
    # This block allows the app to be run directly with `python main.py`.
    # The host '0.0.0.0' is essential for running in a containerized environment,
    # as it allows external connections. The port is determined by the `PORT`
    # environment variable, defaulting to 8080 for local development.
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)